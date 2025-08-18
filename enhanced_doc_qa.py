# IN THIS CODE VLM WORKS ON PAGES WITH VISUALS AND ON OTHER PAGES STANDARD EXTRACTION IS USED



"""
Enhanced Multi-Agent Document Q&A System with Optimized VLM Batch Processing + Metadata Visual Detection

Includes robust fallback: If VLM extraction is weak, falls back to PyPDF, then OCR.
Always tries to report page count for analytics for all file types.
"""

import os
import tempfile
import shutil
from pathlib import Path
from typing import Optional, Dict, Any, List, Tuple
import logging
from contextlib import contextmanager
import uuid
import re
from concurrent.futures import ThreadPoolExecutor, as_completed
import gc
import time
from dotenv import load_dotenv
from typing import TypedDict

load_dotenv()
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
HUGGINGFACE_API_KEY = os.getenv("HUGGINGFACE_API_KEY")

from langgraph.graph import StateGraph, END

class QAState(TypedDict):
    file_path: str
    text: str
    chunks: List[str]
    vectorstore: Optional[object]
    query: str
    answer: str
    retrieved_chunks: List[str]
    extraction_method: str
    word_count: int
    chunk_count: int
    next_action: str


# ====== UTILITY FUNCS FOR PAGE COUNT ======
def get_pdf_page_count(file_path: Path) -> int:
    try:
        import fitz
        doc = fitz.open(str(file_path))
        count = len(doc)
        doc.close()
        return count
    except Exception:
        try:
            from pdf2image import pdfinfo_from_path
            info = pdfinfo_from_path(str(file_path))
            return info['Pages']
        except Exception:
            return 1

def get_docx_page_estimate(text: str) -> int:
    return max(1, len(text) // 1800)

# ====== VISUAL CONTENT DETECTION ======
def has_visual_metadata_pdf(file_path: str) -> bool:
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        total_pages = len(doc)
        pages_with_visuals = []
        
        for page_num in range(total_pages):
            page = doc.load_page(page_num)
            images = page.get_images()
            
            # Check for images in page
            has_images = len(images) > 0
            
            # Alternative check using page dictionary
            text_dict = page.get_text("dict")
            blocks = text_dict.get("blocks", [])
            
            # Look for image blocks
            image_blocks = [b for b in blocks if b.get("type") == 1]  # Type 1 is image
            
            if has_images or image_blocks:
                pages_with_visuals.append(page_num)
                
        doc.close()
        return {
            "has_visuals": bool(pages_with_visuals),
            "visual_pages": pages_with_visuals
        }
    except ImportError:
        return {"has_visuals": False, "visual_pages": []}
    except Exception as e:
        print(f"Error detecting visuals: {e}")
        return {"has_visuals": False, "visual_pages": []}
        # has_actual_images = total_images > 0
        # has_significant_drawings = total_drawings > (total_pages * 25)
        # has_complex_layouts = total_complex_elements > 3
        # if has_actual_images:
        #     return True
        # elif has_significant_drawings:
        #     return True
        # elif has_complex_layouts:
        #     return True
        # else:
        #     return False

def has_visual_metadata_docx(file_path: str) -> bool:
    try:
        from docx import Document
        doc = Document(file_path)
        inline_shapes_count = len(doc.inline_shapes)
        tables_count = len(doc.tables)
        embedded_media_count = 0
        try:
            for rel in doc.part.rels.values():
                if any(media_type in rel.target_ref.lower() for media_type in ["image", "chart", "diagram", "media"]):
                    embedded_media_count += 1
        except:
            pass
        drawing_elements_count = 0
        try:
            xml_content = str(doc._document_part._blob)
            drawing_elements_count = xml_content.count('w:drawing') + xml_content.count('w:pict')
        except:
            pass
        total_visual_elements = inline_shapes_count + tables_count + embedded_media_count + drawing_elements_count
        if total_visual_elements > 0:
            return True
        return False
    except ImportError:
        return False
    except Exception as e:
        return False

def has_visual_content_comprehensive(file_path: str) -> Dict[str, Any]:
    file_path = Path(file_path)
    file_ext = file_path.suffix.lower()
    result = {
        "has_visuals": False,
        "detection_method": "metadata",
        "visual_types": [],
        "confidence": 0.0,
        "details": {}
    }
    try:
        if file_ext == '.pdf':
            # has_visuals = has_visual_metadata_pdf(str(file_path))
            # result["has_visuals"] = has_visuals
            visual_info = has_visual_metadata_pdf(str(file_path))
            result["has_visuals"] = visual_info["has_visuals"]
            result["confidence"] = 0.95 if visual_info["has_visuals"] else 0.9
            if visual_info["has_visuals"]:
                result["visual_types"] = ["images", "tables", "graphics"]
        elif file_ext == '.docx':
            has_visuals = has_visual_metadata_docx(str(file_path))
            result["has_visuals"] = has_visuals
            result["confidence"] = 0.9 if has_visuals else 0.85
            if has_visuals:
                result["visual_types"] = ["inline_shapes", "tables", "embedded_media"]
        elif file_ext == '.txt':
            result["has_visuals"] = False
            result["confidence"] = 1.0
            result["visual_types"] = []
        else:
            result["detection_method"] = "unsupported"
            result["confidence"] = 0.0
    except Exception as e:
        result["detection_method"] = "error"
        result["details"]["error"] = str(e)
    return result

# ====== BATCH VLM PROCESSOR ======
class BatchVLMProcessor:
    def __init__(self, vlm_processor, vlm_model, batch_size=5, max_workers=3):
        self.vlm_processor = vlm_processor
        self.vlm_model = vlm_model
        self.batch_size = batch_size
        self.max_workers = max_workers

    # def process_page_batch(self, page_batch: List[Tuple[int, Any]]) -> List[str]:
    #     print("process_page_batch called with", len(page_batch), "pages")  # <--- Add this
    #     results = []
    #     vlm_prompt = (
    #         "[INST] Extract all visible text, tables (format as markdown tables), "
    #         "and describe any images, figures, or charts from this document page. "
    #         "Preserve the structure and formatting. [/INST]"
    #     )
    #     for page_num, page_img in page_batch:
    #         print(f"Processing page {page_num} in process_page_batch")  # Optional debug
    #         try:
    #             import torch
    #             inputs = self.vlm_processor(
    #                 text=vlm_prompt,
    #                 images=[page_img],
    #                 return_tensors="pt"
    #             ).to(self.vlm_model.device)
    #             with torch.no_grad():
    #                 generation = self.vlm_model.generate(
    #                     **inputs,
    #                     max_new_tokens=512,
    #                     do_sample=True,
    #                     temperature=0.1
    #                 )
    #             page_text = self.vlm_processor.decode(
    #                 generation[0],
    #                 skip_special_tokens=True
    #             )
    #             print(f"VLM output for page {page_num}:", repr(page_text))  # <--- ADD THIS
    #             results.append(f"--- PAGE {page_num} (VLM) ---\n{page_text}\n\n")
    #             del inputs, generation
    #             if hasattr(torch, 'cuda') and torch.cuda.is_available():
    #                 torch.cuda.empty_cache()
    #         except Exception as page_error:
    #             try:
    #                 import pytesseract
    #                 ocr_text = pytesseract.image_to_string(page_img)
    #                 results.append(f"--- PAGE {page_num} (OCR Fallback) ---\n{ocr_text}\n\n")
    #             except:
    #                 results.append(f"--- PAGE {page_num} (FAILED) ---\n[Page processing failed]\n\n")
    #     return results


    def process_page_batch(self, page_batch: List[Tuple[int, Any]]) -> List[str]:
        print("process_page_batch called with", len(page_batch), "pages")
        results = []
        for page_num, page_img in page_batch:
            # If you want to process multiple images, make a list here:
            images = [page_img]  # Or more images if you have them for this page
            num_images = len(images)
            image_tokens = " ".join(["<image>"] * num_images)
            vlm_prompt = (
                f"[INST] {image_tokens} Extract all visible text, tables (format as markdown tables), "
                "and describe any images, figures, or charts from this document page. "
                "Preserve the structure and formatting. [/INST]"
            )
            try:
                import torch
                inputs = self.vlm_processor(
                    text=vlm_prompt,
                    images=images,
                    return_tensors="pt"
                ).to(self.vlm_model.device)
                with torch.no_grad():
                    generation = self.vlm_model.generate(
                        **inputs,
                        max_new_tokens=512,
                        do_sample=True,
                        temperature=0.1
                    )
                page_text = self.vlm_processor.decode(
                    generation[0],
                    skip_special_tokens=True
                )
                if '[/INST]' in page_text:
                    page_text = page_text.split('[/INST]', 1)[1].strip()
                print(f"VLM output for page {page_num}:", repr(page_text))
                results.append(f"--- PAGE {page_num} (VLM) ---\n{page_text}\n\n")
                del inputs, generation
                if hasattr(torch, 'cuda') and torch.cuda.is_available():
                    torch.cuda.empty_cache()
            except Exception as page_error:
                print(f"VLM failed for page {page_num}: {page_error}")
                try:
                    import pytesseract
                    ocr_text = pytesseract.image_to_string(page_img)
                    print(f"OCR output for page {page_num}:", repr(ocr_text))
                    results.append(f"--- PAGE {page_num} (OCR Fallback) ---\n{ocr_text}\n\n")
                except Exception as ocr_error:
                    print(f"OCR failed for page {page_num}: {ocr_error}")
                    results.append(f"--- PAGE {page_num} (FAILED) ---\n[Page processing failed]\n\n")
        return results
    def process_pdf_parallel(self, file_path: Path, visual_pages: List[int] = None) -> Dict[str, Any]:
        try:
            import torch
            import fitz  # PyMuPDF
            import numpy as np
            from PIL import Image

            # Get total pages using PyMuPDF instead of pdf2image
            doc = fitz.open(str(file_path))
            total_pages = len(doc)
            print(f"PDF has {total_pages} pages")

            # First get standard text extraction for all pages
            text_content = {}
            try:
                # Already using PyPDF for text extraction, keep this part
                from pypdf import PdfReader
                reader = PdfReader(str(file_path))
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text() or ""
                    text_content[i] = f"--- PAGE {i+1} (Text) ---\n{page_text}\n\n"
            except Exception as e:
                print(f"Standard text extraction failed: {e}")
                # Create empty placeholders if text extraction fails
                for i in range(total_pages):
                    text_content[i] = f"--- PAGE {i+1} (Text Extraction Failed) ---\n\n"

            # If no visual pages or visual_pages is None, return just the text extraction
            if not visual_pages:
                print("No visual pages to process with VLM")
                extracted_text = ""
                for i in range(total_pages):
                    extracted_text += text_content.get(i, f"--- PAGE {i+1} (Missing) ---\n\n")

                return {
                    "text": extracted_text,
                    "extraction_method": "pypdf",
                    "success": True,
                    "word_count": len(extracted_text.split()),
                    "pages": total_pages,
                    "pages_processed": total_pages,
                    "processing_time": 0,
                }

            print(f"Processing {len(visual_pages)} pages with VLM: {visual_pages}")
            processed_pages = 0

            # Process visual pages in batches
            for start_idx in range(0, len(visual_pages), self.batch_size * self.max_workers):
                batch_indices = visual_pages[start_idx:start_idx + (self.batch_size * self.max_workers)]
                print(f"Processing batch of visual pages: {batch_indices}")

                # Group consecutive pages for efficient conversion
                if not batch_indices:
                    continue

                try:
                    # Create batches for VLM processing
                    page_batches = []
                    current_batch = []

                    # Instead of pdf2image, use PyMuPDF to get page images
                    for page_idx in batch_indices:
                        if 0 <= page_idx < total_pages:  # Ensure page index is valid
                            page = doc.load_page(page_idx)
                            # Render page to an image at a reasonable resolution
                            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better resolution

                            # Convert pixmap to PIL Image
                            img_data = pix.samples
                            img = Image.frombytes("RGB", [pix.width, pix.height], img_data)

                            current_batch.append((page_idx, img))
                            if len(current_batch) >= self.batch_size:
                                page_batches.append(current_batch)
                                current_batch = []

                    # Add any remaining pages
                    if current_batch:
                        page_batches.append(current_batch)

                    # Process batches in parallel
                    with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
                        future_to_batch = {
                            executor.submit(self.process_page_batch, batch): batch
                            for batch in page_batches
                        }

                        for future in as_completed(future_to_batch):
                            batch = future_to_batch[future]
                            try:
                                batch_results = future.result()
                                # Update results for these pages
                                for i, (page_idx, _) in enumerate(batch):
                                    if i < len(batch_results):
                                        text_content[page_idx] = batch_results[i]
                                        processed_pages += 1
                            except Exception as exc:
                                print(f"Batch processing failed: {exc}")
                                for page_idx, _ in batch:
                                    text_content[page_idx] = f"--- PAGE {page_idx+1} (VLM FAILED) ---\n[VLM processing failed]\n\n"

                    # Clean up
                    gc.collect()

                except Exception as chunk_error:
                    print(f"Exception processing visual pages {batch_indices}: {chunk_error}")
                    # Keep text extraction for these pages

            # Close the document
            doc.close()

            # Combine all text in page order
            extracted_text = ""
            for i in range(total_pages):
                extracted_text += text_content.get(i, f"--- PAGE {i+1} (Missing) ---\n\n")

            print(f"process_pdf_parallel: final extraction has {len(extracted_text)} chars")

            if not extracted_text.strip() or len(extracted_text.strip()) < 50:
                # Modified fallback without OCR
                return {
                    "text": "Text extraction failed. Please try another file format.",
                    "extraction_method": "failed",
                    "success": False,
                    "word_count": 0,
                    "pages": total_pages,
                    "pages_processed": 0,
                    "processing_time": 0,
                }

            return {
                "text": extracted_text,
                "extraction_method": "hybrid_vlm_text",
                "success": True,
                "word_count": len(extracted_text.split()),
                "pages": total_pages,
                "pages_processed": total_pages,
                "processing_time": 0,
            }
        except Exception as e:
            return {"success": False, "error": f"Hybrid VLM processing failed: {str(e)}"}

# ====== SMART DOCUMENT PROCESSOR WITH ROBUST FALLBACKS ======
class SmartDocumentProcessor:
    def __init__(self, llm, vlm_processor=None, vlm_model=None, batch_size=5, max_workers=3):
        self.llm = llm
        self.vlm_processor = vlm_processor
        self.vlm_model = vlm_model
        self.temp_dirs = []
        self.batch_processor = BatchVLMProcessor(vlm_processor, vlm_model, batch_size, max_workers) if vlm_processor and vlm_model else None

    @contextmanager
    def temporary_directory(self):
        temp_dir = tempfile.mkdtemp(prefix="doc_qa_")
        self.temp_dirs.append(temp_dir)
        try:
            yield temp_dir
        finally:
            try:
                shutil.rmtree(temp_dir)
                if temp_dir in self.temp_dirs:
                    self.temp_dirs.remove(temp_dir)
            except Exception as e:
                pass

    def extract_text_smart(self, file_path: str) -> Dict[str, Any]:
        file_path = Path(file_path)
        file_ext = file_path.suffix.lower()
        result = {
            "text": "",
            "extraction_method": None,
            "success": False,
            "error": None,
            "file_type": file_ext,
            "has_visual_content": False,
            "visual_detection": {},
            "word_count": 0,
            "pages": 1,
            "processing_time": 0,
        }
        t0 = time.time()
        try:
            visual_analysis = has_visual_content_comprehensive(str(file_path))
            result["has_visual_content"] = visual_analysis["has_visuals"]
            result["visual_detection"] = visual_analysis
            if visual_analysis["has_visuals"] and self.batch_processor:
                extraction = self._extract_with_batch_vlm(file_path)
            else:
                extraction = self._extract_standard(file_path)
            result.update(extraction)
            print("Extracted text from extractor:", result.get("text", "")[:500])
            if "pages" not in result or not result["pages"]:
                if file_ext == ".pdf":
                    result["pages"] = get_pdf_page_count(file_path)
                elif file_ext == ".docx":
                    result["pages"] = get_docx_page_estimate(result.get("text", ""))
                elif file_ext == ".txt":
                    lines = result.get("text", "").count('\n')
                    result["pages"] = max(1, lines // 50)
                else:
                    result["pages"] = 1
            result["processing_time"] = round(time.time() - t0, 2)
            result["success"] = extraction.get("success", False)
        except Exception as e:
            result["error"] = str(e)
        return result

    # def _extract_with_batch_vlm(self, file_path: Path) -> Dict[str, Any]:
    #     try:
    #         if not self.batch_processor:
    #             return {"success": False, "error": "Batch VLM processor not configured"}
    #         if file_path.suffix.lower() == '.pdf':
    #             return self.batch_processor.process_pdf_parallel(file_path)
    #         elif file_path.suffix.lower() == '.docx':
    #             return self._extract_standard(file_path)
    #         else:
    #             return {"success": False, "error": "Batch VLM not supported for this file type"}
    #     except Exception as e:
    #         return self._extract_standard(file_path)
    
    def _extract_with_batch_vlm(self, file_path: Path) -> Dict[str, Any]:
        try:
            if not self.batch_processor:
                return {"success": False, "error": "Batch VLM processor not configured"}

            if file_path.suffix.lower() == '.pdf':
                # Get visual pages information
                visual_info = has_visual_metadata_pdf(str(file_path))
                visual_pages = visual_info.get("visual_pages", [])

                print(f"PDF has {len(visual_pages)} pages with visuals: {visual_pages}")

                # Process with optimized approach
                out = self.batch_processor.process_pdf_parallel(file_path, visual_pages=visual_pages)
                print(f"Extracted text in _extract_with_batch_vlm: {len(out.get('text', ''))} chars")
                return out

            elif file_path.suffix.lower() == '.docx':
                out = self._extract_standard(file_path)
                print(f"Extracted text in _extract_with_batch_vlm: {len(out.get('text', ''))} chars")
                return out

            else:
                return {"success": False, "error": "Batch VLM not supported for this file type"}
            
        except Exception as e:
            out = self._extract_standard(file_path)
            print(f"Extracted text in _extract_with_batch_vlm (exception): {len(out.get('text', ''))} chars")
            return out  

    def _extract_standard(self, file_path: Path) -> Dict[str, Any]:
        file_ext = file_path.suffix.lower()
        try:
            if file_ext == '.pdf':
                out = self._extract_pdf_standard(file_path)
                print("Extracted text in _extract_standard:", out.get("text", "")[:500])  # <--- ADD HERE
                if "pages" not in out or not out.get("pages"):
                    out["pages"] = get_pdf_page_count(file_path)
                return out
            elif file_ext == '.docx':
                out = self._extract_docx_standard(file_path)
                print("Extracted text in _extract_standard:", out.get("text", "")[:500])  # <--- ADD HERE
                if "pages" not in out or not out.get("pages"):
                    out["pages"] = get_docx_page_estimate(out.get("text", ""))
                return out
            elif file_ext == '.txt':
                out = self._extract_text_file(file_path)
                print("Extracted text in _extract_standard:", out.get("text", "")[:500])  # <--- ADD HERE
                if "pages" not in out or not out.get("pages"):
                    lines = out.get("text", "").count('\n')
                    out["pages"] = max(1, lines // 50)
                return out
            else:
                return {"success": False, "error": f"Unsupported file type: {file_ext}"}
        except Exception as e:
            return {"success": False, "error": f"Standard extraction failed: {str(e)}"}

    def _extract_pdf_standard(self, file_path: Path) -> Dict[str, Any]:
        try:
            from pypdf import PdfReader
        except ImportError:
            return {"success": False, "error": "pypdf not installed"}
        try:
            reader = PdfReader(str(file_path))
            text_content = ""
            for page in reader.pages:
                page_text = page.extract_text() or ""
                text_content += page_text + "\n"
            word_count = len(text_content.split())
            pages = len(reader.pages)
            if word_count >= 50:
                return {
                    "text": text_content,
                    "extraction_method": "pypdf",
                    "success": True,
                    "word_count": word_count,
                    "pages": pages,
                }
            else:
                return self._extract_pdf_ocr(file_path)
        except Exception as e:
            return self._extract_pdf_ocr(file_path)

    def _extract_pdf_ocr(self, file_path: Path) -> Dict[str, Any]:
        try:
            import fitz  # PyMuPDF
        except ImportError:
            return {"success": False, "error": "PyMuPDF (fitz) not installed"}
    
        try:
            doc = fitz.open(str(file_path))
            extracted_text = ""

            for i in range(len(doc)):
                page = doc.load_page(i)
                # Try different text extraction methods if one fails
                try:
                    # Try normal text extraction first
                    page_text = page.get_text()
                    if not page_text.strip():
                        # If no text found, try extracting as dict which sometimes gets more text
                        text_dict = page.get_text("dict")
                        blocks = text_dict.get("blocks", [])
                        page_text = "\n".join([block.get("text", "") for block in blocks if "text" in block])
                except:
                    page_text = ""

                extracted_text += f"--- PAGE {i+1} (PyMuPDF) ---\n{page_text}\n\n"

            doc.close()
            return {
                "text": extracted_text,
                "extraction_method": "pymupdf",
                "success": True,
                "word_count": len(extracted_text.split()),
                "pages": len(doc),
            }
        except Exception as e:
            return {"success": False, "error": f"PyMuPDF extraction failed: {str(e)}"}

    def _extract_docx_standard(self, file_path: Path) -> Dict[str, Any]:
        try:
            from docx import Document as DocxDoc
            doc = DocxDoc(str(file_path))
            parts = []
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
            for tbl in doc.tables:
                for row in tbl.rows:
                    cells = [c.text.strip() for c in row.cells]
                    parts.append("\t".join(cells))
            text_content = "\n".join(parts).strip()
            word_count = len(text_content.split())
            pages = get_docx_page_estimate(text_content)
            if word_count >= 20:
                return {
                    "text": text_content,
                    "extraction_method": "python-docx+tables",
                    "success": True,
                    "word_count": word_count,
                    "pages": pages,
                }
            import docx2txt
            text_content = docx2txt.process(str(file_path))
            word_count = len(text_content.split())
            pages = get_docx_page_estimate(text_content)
            return {
                "text": text_content,
                "extraction_method": "docx2txt",
                "success": True,
                "word_count": word_count,
                "pages": pages,
            }
        except ImportError:
            try:
                from langchain_community.document_loaders import Docx2txtLoader
                loader = Docx2txtLoader(str(file_path))
                documents = loader.load()
                text_content = "\n".join(doc.page_content for doc in documents)
                word_count = len(text_content.split())
                pages = get_docx_page_estimate(text_content)
                return {
                    "text": text_content,
                    "extraction_method": "langchain_docx",
                    "success": True,
                    "word_count": word_count,
                    "pages": pages,
                }
            except ImportError:
                return {"success": False, "error": "DOCX deps missing (docx/docx2txt)"}
        except Exception as e:
            return {"success": False, "error": f"DOCX extraction failed: {str(e)}"}

    def _extract_text_file(self, file_path: Path) -> Dict[str, Any]:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text_content = file.read()
            num_lines = text_content.count('\n')
            num_pages = max(1, num_lines // 50)
            return {
                "text": text_content,
                "extraction_method": "direct",
                "success": True,
                "word_count": len(text_content.split()),
                "pages": num_pages,
            }
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text_content = file.read()
                num_lines = text_content.count('\n')
                num_pages = max(1, num_lines // 50)
                return {
                    "text": text_content,
                    "extraction_method": "direct",
                    "success": True,
                    "word_count": len(text_content.split()),
                    "pages": num_pages,
                }
            except Exception as e:
                return {"success": False, "error": f"Text file extraction failed: {str(e)}"}
        except Exception as e:
            return {"success": False, "error": f"Text file extraction failed: {str(e)}"}

    def cleanup(self):
        for temp_dir in self.temp_dirs[:]:
            try:
                if os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir)
                    self.temp_dirs.remove(temp_dir)
            except Exception as e:
                pass

# ====== Paste your InMemoryVectorStore, HallucinationResistantAnswerer, agent functions, and main workflow below ======
class InMemoryVectorStore:
    """In-memory vector store with session isolation."""

    def __init__(self, session_id: Optional[str] = None):
        self.session_id = session_id or str(uuid.uuid4())
        self.vectorstore = None
        self.chunks = None

    def create_vectorstore(self, text: str, chunk_size: int = 1000,
                      chunk_overlap: int = 200) -> Dict[str, Any]:
        """Create in-memory vector store with dynamic parameters."""
        try:
            from langchain.text_splitter import RecursiveCharacterTextSplitter
            from langchain.vectorstores import FAISS
            from langchain_huggingface import HuggingFaceEmbeddings
        except ImportError as e:
            return {"success": False, "error": f"Vector dependencies missing: {e}"}

        try:
            # Dynamic chunk size based on document length, but with higher overlap
            doc_length = len(text)

            # Increase overlap to 50% of chunk size for better continuity
            if doc_length < 2000:
                chunk_size = max(50, doc_length // 2)
                chunk_overlap = max(25, chunk_size // 2)  # 50% overlap
            elif doc_length > 50000:
                chunk_size = 1500
                chunk_overlap = 750  # 50% overlap
            else:
                chunk_size = 800  # Smaller chunks
                chunk_overlap = 400  # 50% overlap

            logger.info(f"Using chunk_size={chunk_size}, chunk_overlap={chunk_overlap}")
        # try:
        #     from langchain.text_splitter import RecursiveCharacterTextSplitter
        #     from langchain.vectorstores import FAISS  # Changed from Chroma to FAISS
        #     from langchain_huggingface import HuggingFaceEmbeddings
        # except ImportError as e:
        #     return {"success": False, "error": f"Vector dependencies missing: {e}"}

        # try:
        #     # Dynamic chunk size based on document length
        #     doc_length = len(text)
        #     chunk_size = 500
        #     chunk_overlap = 100
        #     if doc_length < 2000:
        #         chunk_size = max(50, doc_length // 2)
        #         chunk_overlap = max(10, chunk_size // 10)
        #     elif doc_length > 50000:
        #         chunk_size = 1500
        #         chunk_overlap = 200

        #     logger.info(f"Using chunk_size={chunk_size}, chunk_overlap={chunk_overlap}")
            def _clean_text_for_split(t: str) -> str:
                import re
                t = t.replace('|', '\t')
                t = re.sub(r'[ \t]+', ' ', t)
                t = re.sub(r'\n{3,}', '\n\n', t)
                return t.strip()

            text = _clean_text_for_split(text)
            if len(text) < 50:
                return {"success": False, "error": "No usable text after extraction/cleaning."}
    
            # If short doc, make one bigger chunk
            if len(text) < 500:
                chunk_size = max(100, len(text))
                chunk_overlap = 0       
            # Create chunks
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                separators=["\n\n", "\n", ". ", " ", ""]
            )

            self.chunks = splitter.split_text(text)

            logger.info(f"Text splitter produced {len(self.chunks)} chunks.")
            if not self.chunks or len(self.chunks) == 0:
                logger.error("No text chunks to create vector store! Extraction may have failed or document is empty.")
                return {"success": False, "error": "No data found for embedding."}

            # Create embeddings and in-memory vectorstore
            embeddings = HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-MiniLM-L6-v2"
            )

            # Create FAISS index from documents
            self.vectorstore = FAISS.from_texts(
                self.chunks, 
                embeddings
            )

            return {
                "success": True,
                "chunk_count": len(self.chunks),
                "chunk_size": chunk_size,
                "chunk_overlap": chunk_overlap
            }

        except Exception as e:
            logger.error(f"Vector store creation failed: {e}")
            return {"success": False, "error": str(e)}

    def retrieve_chunks(self, query: str, k: int = 8) -> List[str]:
        """Retrieve relevant chunks for the query."""
        if not self.vectorstore:
            return []

        try:
            # FAISS has a similar retriever interface to Chroma
            retriever = self.vectorstore.as_retriever(
                search_type='mmr', #Maximum Marginal Relevance
                search_kwargs={"k": min(k, len(self.chunks) if self.chunks else k),
                "fetch_k": min(k*3, len(self.chunks) if self.chunks else k*3),
                "lambda_mult": 0.7  # Balance between relevance and diversity
                }
            )

            docs = retriever.get_relevant_documents(query)
            return [doc.page_content for doc in docs]

        except Exception as e:
            logger.error(f"Chunk retrieval failed: {e}")
            return []

class HallucinationResistantAnswerer:
    """Answerer with advanced hallucination prevention."""

    def __init__(self, llm):
        self.llm = llm

    def generate_answer(self, query: str, context_chunks: List[str]) -> Dict[str, Any]:
        """Generate answer with hallucination prevention techniques."""
        if not context_chunks:
            return {
                "answer": "I don't have enough information to answer this question based on the provided document.",
                "confidence": 0.0,
                "sources_used": 0
            }

        # Check if LLM is available
        if not self.llm:
            return {
                "answer": "AI model not available. Please configure your API keys to enable AI-powered answers. For now, here are the relevant document chunks:\n\n" + "\n\n---\n\n".join(context_chunks[:3]),
                "confidence": 0.1,
                "sources_used": len(context_chunks)
            }

        # Create numbered context
        context = "\n\n".join(f"[Source {i+1}]: {chunk}" for i, chunk in enumerate(context_chunks))

        # Advanced anti-hallucination prompt
        enhanced_prompt = f"""You are a precise document analysis assistant. Your task is to answer questions using ONLY the provided context.

CRITICAL RULES:
1. Use ONLY information from the provided sources below
2. If the sources don't contain enough information, explicitly state: "The provided sources don't contain sufficient information to answer this question"
3. Give best answer possible
4. Never add information from your general knowledge
5. If uncertain about any part of your answer, state your uncertainty clearly
6. Provide specific quotes when relevant
7. Present information in a well-structured format.
8. For schedules, timelines, or lists, ALWAYS check if there are any missing weeks, days, or entries.
Don't write the word "Sources" like "Sources 1", "Sources 2", etc. please in your response.


SOURCES:
{context}

QUESTION: {query}

ANALYSIS AND ANSWER (cite sources and be precise):"""

        try:
            result = self.llm.invoke(enhanced_prompt)
            answer = result.content if hasattr(result, 'content') else str(result)

            # Confidence estimation
            confidence = self._estimate_confidence(answer, context_chunks, query)

            return {
                "answer": answer,
                "confidence": confidence,
                "sources_used": len(context_chunks),
                "context_length": len(context)
            }

        except Exception as e:
            logger.error(f"Answer generation failed: {e}")
            return {
                "answer": f"Error generating answer: {str(e)}",
                "confidence": 0.0,
                "sources_used": 0
            }

    def _estimate_confidence(self, answer: str, context_chunks: List[str], query: str) -> float:
        """Estimate answer confidence based on multiple factors."""
        answer_lower = answer.lower()

        # Low confidence indicators
        if any(phrase in answer_lower for phrase in [
            "don't have enough", "not enough information", "insufficient information",
            "unclear", "uncertain", "cannot determine", "not specified"
        ]):
            return 0.1

        # Medium-low confidence
        if any(phrase in answer_lower for phrase in [
            "appears to", "seems to", "might be", "could be", "possibly"
        ]):
            return 0.4

        # Check source citation (good indicator of grounded response)
        has_citations = any(phrase in answer_lower for phrase in [
            "source", "according to", "states that", "mentions", "indicates"
        ])

        # Check context overlap
        context_text = " ".join(context_chunks).lower()
        query_lower = query.lower()

        # Extract key terms from query and answer
        query_terms = set(query_lower.split()) - {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by'}
        answer_terms = set(answer_lower.split()) - {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by'}
        context_terms = set(context_text.split()) - {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by'}

        # Calculate overlaps
        answer_context_overlap = len(answer_terms.intersection(context_terms)) / len(answer_terms) if answer_terms else 0
        query_answer_relevance = len(query_terms.intersection(answer_terms)) / len(query_terms) if query_terms else 0

        # Base confidence calculation
        base_confidence = (answer_context_overlap * 0.6) + (query_answer_relevance * 0.4)

        # Boost for citations
        if has_citations:
            base_confidence = min(0.95, base_confidence * 1.2)

        # Ensure reasonable bounds
        return max(0.2, min(0.95, base_confidence))

# LangGraph Agent Node Functions (updated with enhanced metadata detection)
def document_extractor(state: QAState) -> QAState:
    """Agent 1: Extract text from document with metadata-based visual detection and optimized batch VLM processing."""
    logger.info("=== AGENT 1: Document Extractor (Enhanced) ===")
    
    try:
        # Initialize processor with your actual models
        processor = SmartDocumentProcessor(
            llm=llm,  # Replace with your LLM
            vlm_processor=vlm_processor,  # Replace with your VLM processor
            vlm_model=vlm_model,  # Replace with your VLM model
            batch_size=5,  # Adjust based on your memory constraints
            max_workers=3   # Adjust based on your CPU/GPU resources
        )

        extraction_result = processor.extract_text_smart(state["file_path"])

        if extraction_result["success"]:
            state["text"] = extraction_result["text"]
            state["extraction_method"] = extraction_result["extraction_method"]
            state["word_count"] = extraction_result["word_count"]

            # QUALITY GUARD + AUTO-RECOVERY for weak extraction
        txt = (state["text"] or "").strip()
        if len(txt) < 150:
            logger.warning("⚠️ Extracted text too small; attempting recovery (VLM/OCR)...")
            from pathlib import Path
            recovered = ""
        
            # Prefer VLM for PDFs if batch processor is available and visuals were detected
            visual_info = extraction_result.get("visual_detection", {}) or {}
            try:
                if Path(state["file_path"]).suffix.lower() == ".pdf":
                    if processor.batch_processor and visual_info.get("has_visuals", False):
                        alt = processor._extract_with_batch_vlm(Path(state["file_path"]))
                        if alt.get("success"):
                            recovered = (alt.get("text") or "").strip()
        
                    # If still weak, force OCR fallback
                    if len(recovered) < 150:
                        alt = processor._extract_pdf_ocr(Path(state["file_path"]))
                        if alt.get("success"):
                            recovered = (alt.get("text") or "").strip()
            except Exception as rec_e:
                logger.error(f"Recovery attempt failed: {rec_e}")
        
            if len(recovered) >= 150:
                logger.info("✅ Recovery succeeded; proceeding with recovered text.")
                state["text"] = recovered
                state["extraction_method"] = f"{state['extraction_method']}_recovered"
                state["word_count"] = len(recovered.split())
            else:
                logger.error("❌ Could not recover sufficient text. Aborting before vectorization.")
                state["text"] = ""
                state["extraction_method"] = "failed"
                state["word_count"] = 0
                processor.cleanup()
                return state
        

            logger.info(f"✅ Extraction successful: {extraction_result['word_count']} words")
            logger.info(f"📄 Method: {extraction_result['extraction_method']}")

            # Enhanced logging for visual detection
            visual_info = extraction_result.get("visual_detection", {})
            if visual_info.get("has_visuals"):
                logger.info(        
                    f"🎯 Visual content: {visual_info.get('visual_types', [])} "
                    f"(confidence: {visual_info.get('confidence', 0):.2f})"
                )
            else:
                logger.info(
                    f"📄 Text-only document (confidence: {visual_info.get('confidence', 0):.2f})"
                )

            if extraction_result["extraction_method"] == "vlm_batch":
                logger.info(
                    f"🚀 Batch VLM processing completed: {extraction_result.get('pages_processed', 0)} pages"
                )

        else:
            logger.error(f"❌ Extraction failed: {extraction_result['error']}")
            state["text"] = ""
            state["extraction_method"] = "failed"
            state["word_count"] = 0


        processor.cleanup()

    except Exception as e:
        logger.error(f"Document Extractor error: {e}")
        state["text"] = ""
        state["extraction_method"] = "error"
        state["word_count"] = 0

    return state

def synthesizer(state: QAState) -> QAState:
    logger.info("=== AGENT 2: Synthesizer ===")
    try:
        # Put this guard before any vectorstore creation
        if not state["text"].strip():
            logger.error("❌ Document extraction produced no usable text; skipping vectorization.")
            state["vectorstore"] = None
            state["chunk_count"] = 0
            return state

        vectorstore = InMemoryVectorStore()
        vs_result = vectorstore.create_vectorstore(state["text"])

        if vs_result["success"]:
            state["vectorstore"] = vectorstore
            state["chunk_count"] = vs_result["chunk_count"]
            logger.info(f"✅ Vector store created with {vs_result['chunk_count']} chunks")
        else:
            logger.error(f"❌ Vector store creation failed: {vs_result['error']}")
            state["vectorstore"] = None
            state["chunk_count"] = 0

    except Exception as e:
        logger.error(f"Synthesizer error: {e}")
        state["vectorstore"] = None
        state["chunk_count"] = 0

    return state


def query_answerer(state: QAState) -> QAState:
    logger.info("=== AGENT 3: Query Answerer ===")
    try:
        # Put this check first
        if not state["vectorstore"]:
            state["answer"] = (
                "Vector index was not created because extraction didn’t yield enough usable text. "
                "For PDFs, enable VLM and/or install OCR dependencies (tesseract-ocr, poppler-utils, "
                "pdf2image, pytesseract). For DOCX tables, use the table-aware extractor. Then retry."
            )
            state["retrieved_chunks"] = []
            return state

        # Normal path continues here
        chunks = state["vectorstore"].retrieve_chunks(state["query"], k=6)
        state["retrieved_chunks"] = chunks

        answerer = HallucinationResistantAnswerer(llm=llm)
        answer_result = answerer.generate_answer(state["query"], chunks)

        state["answer"] = answer_result["answer"]
        logger.info(f"✅ Answer generated using {len(chunks)} chunks")
        logger.info(f"🎯 Confidence: {answer_result['confidence']:.2f}")

    except Exception as e:
        logger.error(f"Query Answerer error: {e}")
        state["answer"] = f"Error answering query: {str(e)}"
        state["retrieved_chunks"] = []

    return state

def ask_follow_up(state: QAState) -> QAState:
    """Agent 4: Handle follow-up questions and control loop."""
    logger.info("=== AGENT 4: Ask Follow Up ===")

    # Display current answer
    print(f"\n{'='*60}")
    print("💡 ANSWER:")
    print(f"{'='*60}")
    print(state["answer"])
    print(f"{'='*60}")
    print(f"📚 Sources used: {len(state.get('retrieved_chunks', []))}")
    print(f"{'='*60}")

    # Ask for follow-up
    follow_up = input("\n❓ Do you have another question? (yes/no): ").strip().lower()
    
    if follow_up in ['yes', 'y']:
        new_query = input("🔍 Your question: ").strip()
        if new_query:
            state["query"] = new_query
            state["next_action"] = "continue"
            return state
        else:
            print("No question provided.")

    # End the conversation
    print("👋 Thank you! Exiting...")
    state["next_action"] = "end"
    return state

def should_continue(state: QAState) -> str:
    """Determine if we should continue the Q&A loop or end."""
    return state.get("next_action", "end")

# Main LangGraph Workflow
def create_multiagent_workflow():
    """Create the LangGraph workflow."""
    workflow = StateGraph(QAState)

    # Add nodes
    workflow.add_node("document_extractor", document_extractor)
    workflow.add_node("synthesizer", synthesizer)
    workflow.add_node("query_answerer", query_answerer)
    workflow.add_node("ask_follow_up", ask_follow_up)

    # Add edges
    workflow.add_edge("document_extractor", "synthesizer")
    workflow.add_edge("synthesizer", "query_answerer")
    workflow.add_edge("query_answerer", "ask_follow_up")

    # Conditional edge for looping
    workflow.add_conditional_edges(
        "ask_follow_up",
        should_continue,
        {
            "continue": "query_answerer",
            "end": END
        }
    )

    # Set entry point
    workflow.set_entry_point("document_extractor")

    return workflow.compile()

def run_document_qa_system(file_path: str, llm=None, vlm_processor=None, vlm_model=None, 
                          batch_size=5, max_workers=3):
    """Main function to run the enhanced document Q&A system with metadata-based visual detection."""
    print(f"🚀 Enhanced Multi-Agent Document Q&A System")
    print(f"📄 Processing: {file_path}")
    print(f"🔍 Features: Metadata visual detection + Batch VLM processing")
    print(f"⚙️ Config: Batch size={batch_size}, Max workers={max_workers}")

    # Validate file exists
    if not os.path.exists(file_path):
        print(f"❌ File not found: {file_path}")
        return

    # Get initial query
    initial_query = input("\n🔍 What would you like to know about this document? ").strip()
    if not initial_query:
        print("❌ No initial query provided.")
        return

    # Initialize state
    initial_state = QAState(
        file_path=file_path,
        text="",
        chunks=[],
        vectorstore=None,
        query=initial_query,
        answer="",
        retrieved_chunks=[],
        extraction_method="",
        word_count=0,
        chunk_count=0,
        next_action="continue"
    )

    # Create and run workflow
    try:
        workflow = create_multiagent_workflow()
        print("\n🔄 Starting enhanced workflow with metadata detection...")
        
        final_state = workflow.invoke(initial_state)
        
        print("\n✅ Enhanced workflow completed successfully!")
        print(f"📊 Final stats: {final_state.get('word_count', 0)} words, {final_state.get('chunk_count', 0)} chunks")
        
    except Exception as e:
        logger.error(f"Workflow execution failed: {e}")
        print(f"❌ Error: {e}")

# Entry point
if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        # ... print usage ...
        sys.exit(1)

    file_path = sys.argv[1]
    batch_size = int(sys.argv[2]) if len(sys.argv) > 2 else 5
    max_workers = int(sys.argv[3]) if len(sys.argv) > 3 else 3

    # Initialize your actual models here
    from langchain_google_genai import ChatGoogleGenerativeAI
    from transformers import AutoProcessor, AutoModelForVision2Seq

    llm = ChatGoogleGenerativeAI(model='gemini-2.0-flash', google_api_key=GEMINI_API_KEY)
    vlm_processor = AutoProcessor.from_pretrained("HuggingFaceTB/SmolVLM-256M-Instruct",token=HUGGINGFACE_API_KEY)
    vlm_model = AutoModelForVision2Seq.from_pretrained("HuggingFaceTB/SmolVLM-256M-Instruct",token=HUGGINGFACE_API_KEY)

    # Expose models globally for agent functions
    globals()["llm"] = llm
    globals()["vlm_processor"] = vlm_processor
    globals()["vlm_model"] = vlm_model

    run_document_qa_system(file_path, llm=llm, vlm_processor=vlm_processor, vlm_model=vlm_model, 
        batch_size=batch_size, max_workers=max_workers)
    
